import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import fitz  # PyMuPDF
import os
import sys

# Paths
base_dir = r"c:\xampp\htdocs\uiri-ims\documentation"
db_design_dir = os.path.join(base_dir, "database_design")
sys_design_dir = os.path.join(base_dir, "system_design")
presentation_dir = os.path.join(base_dir, "presentation")

erd_pdf = os.path.join(base_dir, "ERD_IMS(UIRI).pdf")
use_cases_pdf = os.path.join(presentation_dir, "Use-Cases(IMS-DBMS).pdf")

erd_png = os.path.join(base_dir, "ERD_temp.png")
use_cases_png = os.path.join(base_dir, "Use_Cases_temp.png")

output_docx = os.path.join(db_design_dir, "UIRI_IMS_Database_Design_v2_Professional.docx")

# 1. Convert PDFs to PNG
print("Converting PDFs to PNG...")
try:
    if os.path.exists(erd_pdf):
        doc = fitz.open(erd_pdf)
        page = doc.load_page(0)
        pix = page.get_pixmap(dpi=300)
        pix.save(erd_png)
        print("ERD PNG generated.")
    else:
        print(f"Warning: {erd_pdf} not found.")

    if os.path.exists(use_cases_pdf):
        doc = fitz.open(use_cases_pdf)
        page = doc.load_page(0)
        pix = page.get_pixmap(dpi=300)
        pix.save(use_cases_png)
        print("Use Cases PNG generated.")
    else:
        print(f"Warning: {use_cases_pdf} not found.")
except Exception as e:
    print(f"Error converting PDFs: {e}")

# 2. Generate Word Document
print("Generating Word Document...")
document = docx.Document()

# Styles
styles = document.styles

# Title Style
title_style = styles.add_style('ProfessionalTitle', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
title_font = title_style.font
title_font.name = 'Arial'
title_font.size = Pt(28)
title_font.bold = True
title_font.color.rgb = RGBColor(15, 36, 62)
title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_style.paragraph_format.space_after = Pt(24)

# Heading 1 Style
h1_style = styles['Heading 1']
h1_font = h1_style.font
h1_font.name = 'Arial'
h1_font.size = Pt(18)
h1_font.bold = True
h1_font.color.rgb = RGBColor(31, 73, 125)
h1_style.paragraph_format.space_before = Pt(24)
h1_style.paragraph_format.space_after = Pt(12)
h1_style.paragraph_format.keep_with_next = True

# Heading 2 Style
h2_style = styles['Heading 2']
h2_font = h2_style.font
h2_font.name = 'Arial'
h2_font.size = Pt(14)
h2_font.bold = True
h2_font.color.rgb = RGBColor(79, 129, 189)
h2_style.paragraph_format.space_before = Pt(18)
h2_style.paragraph_format.space_after = Pt(6)

# Normal text
normal_style = styles['Normal']
normal_font = normal_style.font
normal_font.name = 'Calibri'
normal_font.size = Pt(11)
normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
normal_style.paragraph_format.space_after = Pt(10)

# Cover Page
document.add_paragraph('UIRI Inventory Management System', style='ProfessionalTitle')
document.add_paragraph('Database Design & System Architecture', style='ProfessionalTitle').runs[0].font.size = Pt(22)
document.add_paragraph('Version 2.0 - Professional Revision', style='ProfessionalTitle').runs[0].font.size = Pt(14)
document.add_paragraph('\n\n\n')

# Document Info
p = document.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Uganda Industrial Research Institute (UIRI)\nPrepared by: Systems Development Team\nDate: July 2026')
run.font.size = Pt(12)
run.font.italic = True
document.add_page_break()

# 1. Introduction
document.add_heading('1. Introduction', level=1)
document.add_paragraph('This document outlines the revised database design and system architecture for the UIRI Inventory Management System (IMS). It integrates the Data Dictionary, logical Entity-Relationship Diagram (ERD), and System Use Cases to provide a comprehensive view of the system\'s data structures and interactions.')
document.add_paragraph('The database is designed to handle multi-branch inventory tracking, procurement workflows, stock movements, and strict role-based access control (RBAC).')

# 2. System Use Cases
document.add_heading('2. System Architecture & Use Cases', level=1)
document.add_paragraph('The UIRI IMS caters to various actors including Administrators, Storekeepers, Department Heads, and standard Staff. The following diagram illustrates the core use cases and interactions within the system:')

if os.path.exists(use_cases_png):
    try:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(use_cases_png, width=Inches(6.0))
        p_caption = document.add_paragraph('Figure 1: UIRI IMS System Use Cases', style='Normal')
        p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_caption.runs[0].font.italic = True
    except Exception as e:
        document.add_paragraph(f'[Error inserting Use Cases Image: {e}]')
else:
    document.add_paragraph('[System Use Cases Image Not Found - Placeholder]')

# 3. Database Logical Design (ERD)
document.add_heading('3. Logical Database Design (ERD)', level=1)
document.add_paragraph('The logical ERD represents the conceptual structure of the UIRI IMS database. It is segmented into distinct domains: Organization & Access, Catalog & Inventory Master, Stock Operations, Procurement, and Equipment Maintenance.')

if os.path.exists(erd_png):
    try:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(erd_png, width=Inches(6.5))
        p_caption = document.add_paragraph('Figure 2: Conceptual Entity-Relationship Diagram', style='Normal')
        p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_caption.runs[0].font.italic = True
    except Exception as e:
        document.add_paragraph(f'[Error inserting ERD Image: {e}]')
else:
    document.add_paragraph('[Logical ERD Image Not Found - Placeholder]')

# 4. Data Dictionary (Revised Understanding)
document.add_heading('4. Data Dictionary (Core Entities)', level=1)
document.add_paragraph('The following section details the core entities within the system, replacing the legacy Data Dictionary with an updated understanding of the schema.')

def add_entity_table(doc, entity_name, description, columns):
    doc.add_heading(entity_name, level=2)
    doc.add_paragraph(description)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(['Column Name', 'Data Type', 'Constraints', 'Description']):
        hdr_cells[i].text = text
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        
    for col in columns:
        row_cells = table.add_row().cells
        row_cells[0].text = col[0]
        row_cells[1].text = col[1]
        row_cells[2].text = col[2]
        row_cells[3].text = col[3]

add_entity_table(document, '4.1 Users & Authentication (users)', 
                 'Manages system access and role assignment.',
                 [
                     ('id', 'INT', 'PK, AI', 'Unique identifier'),
                     ('username', 'VARCHAR(50)', 'UNIQUE, NOT NULL', 'Login name'),
                     ('role_id', 'INT', 'FK', 'References roles(id)'),
                     ('branch_id', 'INT', 'FK', 'References branches(id)'),
                     ('password_hash', 'VARCHAR(255)', 'NOT NULL', 'Bcrypt hashed password')
                 ])

document.add_paragraph('\n')

add_entity_table(document, '4.2 Inventory Master (items)', 
                 'Core catalog of all items tracked across branches.',
                 [
                     ('id', 'INT', 'PK, AI', 'Unique identifier'),
                     ('item_code', 'VARCHAR(50)', 'UNIQUE, NOT NULL', 'SKU or Barcode'),
                     ('name', 'VARCHAR(100)', 'NOT NULL', 'Item name'),
                     ('category_id', 'INT', 'FK', 'References categories(id)'),
                     ('minimum_stock', 'INT', 'DEFAULT 0', 'Reorder threshold')
                 ])

document.add_paragraph('\n')

add_entity_table(document, '4.3 Stock Transactions (stock_transactions)', 
                 'Immutable ledger of all stock movements (IN/OUT/ADJ).',
                 [
                     ('id', 'INT', 'PK, AI', 'Unique identifier'),
                     ('item_id', 'INT', 'FK, NOT NULL', 'References items(id)'),
                     ('branch_id', 'INT', 'FK, NOT NULL', 'References branches(id)'),
                     ('transaction_type', 'ENUM', 'NOT NULL', "'IN', 'OUT', or 'ADJUSTMENT'"),
                     ('quantity', 'DECIMAL', 'NOT NULL', 'Amount moved'),
                     ('user_id', 'INT', 'FK, NOT NULL', 'User who performed action')
                 ])

# 5. Conclusion
document.add_heading('5. Conclusion', level=1)
document.add_paragraph('The revised architecture implements strong data integrity via foreign key constraints, optimizes reporting through structured categorical entities, and ensures accountability via the comprehensive stock_transactions ledger and audit logs.')

# Save Document
document.save(output_docx)
print(f"\nSuccessfully created professional Word Document: {output_docx}")

# Cleanup temp images
if os.path.exists(erd_png): os.remove(erd_png)
if os.path.exists(use_cases_png): os.remove(use_cases_png)
